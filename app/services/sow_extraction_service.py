"""SOW requirement extraction -- reads an already-uploaded Statement of Work
(PDF or Word) for a project and uses this app's own LLM provider chain (same
Azure OpenAI/Gemini/Claude failover as resume_skill_service.py) to pull out
the roles, skills, and competency areas the engagement actually calls for.

Grounding rules, same posture as resume_skill_service.py:
  - Roles: extracted only from staffing/resourcing text that's actually in
    the SOW (e.g. a "Resources" table listing designations and headcounts),
    never invented. Each extracted role is cross-checked against the org's
    real designation list (01_Employee_Details_clean.csv) and flagged
    matches_real_designation so a "Snr. Oversight" (an engagement-level
    label, not an HR designation) reads differently from a "Solutions
    Consultant" (a real designation) -- this is the accuracy signal the
    "100% accurate" ask actually cashes out to: never claim a role is a real
    designation unless it genuinely is one.
  - Skills: only technologies/tools/methodologies literally named in the
    SOW's scope/approach/deliverables text (e.g. "Power BI", "Microsoft
    Fabric", "logistic regression") -- never inferred from a vague
    description of the work.
  - Competency areas: short, SOW-grounded phrases describing the kind of
    expertise or behaviour the engagement will require (e.g. "client
    stakeholder workshops", "technical documentation"). Deliberately NOT
    mapped onto the org's internal HR competency-review statements (see
    06_Competency_Details_clean.csv) -- a SOW's staffing requirements and
    an employee's internal appraisal framework are two different things,
    and force-fitting one onto the other would fabricate a link that
    doesn't exist in the source data.

Purely a display/proof feature for now -- not wired into recommendation
scoring, the budget step, or anything else, per explicit instruction.
"""
import io
import json
from pathlib import Path

import pandas as pd
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ai import llm
from app.core.config import APP_STATE_DIR, TRANSFORMED_DIR
from app.services.project_sow_service import get_sow_file_path

EXTRACTIONS_JSON_PATH = APP_STATE_DIR / "sow_extractions.json"
EMPLOYEES_CSV_PATH = TRANSFORMED_DIR / "01_Employee_Details_clean.csv"

MAX_SOW_CHARS = 20000


class SowExtractionError(Exception):
    pass


def _iter_block_items(document):
    """Walk a docx body in document order, yielding Paragraph and Table
    objects as they actually appear -- python-docx's own `.paragraphs` and
    `.tables` collections lose that interleaving, which matters here since a
    heading like "Resources:" and the staffing table beneath it need to stay
    read in the same order a human reads the document."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _docx_table_text(table: Table) -> str:
    lines = []
    for row in table.rows:
        cells = [c.text.strip() for c in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_text(filename: str, content: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        from pypdf import PdfReader
        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception as exc:
            raise SowExtractionError(f"Could not read this PDF: {exc}") from exc
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif ext == "docx":
        import docx
        try:
            doc = docx.Document(io.BytesIO(content))
        except Exception as exc:
            raise SowExtractionError(f"Could not read this Word document: {exc}") from exc
        parts = []
        for block in _iter_block_items(doc):
            if isinstance(block, Paragraph):
                if block.text.strip():
                    parts.append(block.text.strip())
            else:
                table_text = _docx_table_text(block)
                if table_text:
                    parts.append(table_text)
        text = "\n".join(parts)
    else:
        raise SowExtractionError(f"Unsupported file type '.{ext}' -- upload a PDF or Word (.docx) SOW.")

    text = text.strip()
    if not text:
        raise SowExtractionError("Could not extract any real text from this file -- it may be a scanned image without a real text layer.")
    return text[:MAX_SOW_CHARS]


def _real_designations() -> list[str]:
    df = pd.read_csv(EMPLOYEES_CSV_PATH, dtype=str)
    return sorted(df["job_name"].dropna().str.strip().unique().tolist())


_EXTRACTION_PROMPT_TEMPLATE = """You are extracting real staffing and skill requirements from a real client Statement of Work (SOW) for an internal resourcing tool. Extract ONLY what the document actually says -- never invent a role, skill, or requirement that isn't genuinely present in the text below.

PART 1 -- ROLES REQUIRED
Find the staffing/resourcing section (often a table listing job titles/designations, sometimes with a headcount like "2x Software Engineers" or a named person like "Solutions Consultant (Jane Doe)"). For each distinct role, extract:
- role_text: the exact role/designation text as written (without the count prefix or the named person)
- count: how many of this role (default 1 if not specified, e.g. "2x Software Engineers" -> count 2)
- named_person: the specific person's name if one is given for this role, else null

PART 2 -- SKILLS REQUIRED
Extract only specific technologies, tools, platforms, or methodologies that are LITERALLY NAMED somewhere in the scope of work, approach, or deliverables text (e.g. "Power BI", "Microsoft Fabric", "logistic regression", "Python"). Do not infer a skill from a vague description of work that doesn't actually name a technology.

PART 3 -- COMPETENCY AREAS REQUIRED
Extract short phrases (3-6 words each) describing the kind of expertise or behaviour this engagement will require, each grounded in specific text in the document (e.g. "client stakeholder workshops", "technical documentation production", "executive-level reporting"). Do not use generic filler -- each phrase must correspond to something the document actually describes the team doing.

PART 4 -- LIGHT CONTEXT (only if genuinely present in the text)
- client_name: the client company name, or null
- project_reference: JMAN's internal project/engagement reference code if one appears, or null
- engagement_duration: the stated timeline/duration (e.g. "8 weeks"), or null
- scope_summary: one plain sentence summarizing what the engagement is building/doing, based only on the text

Respond with ONLY a JSON object, no other text, no markdown fences:
{{
  "roles_required": [{{"role_text": "...", "count": 1, "named_person": null}}],
  "skills_required": ["..."],
  "competency_areas_required": ["..."],
  "client_name": null,
  "project_reference": null,
  "engagement_duration": null,
  "scope_summary": null
}}

SOW text:
---
{sow_text}
---
"""


def _build_extraction_prompt(sow_text: str) -> str:
    return _EXTRACTION_PROMPT_TEMPLATE.format(sow_text=sow_text)


def _call_llm_for_extraction(sow_text: str) -> dict:
    providers = llm.get_providers()
    if not providers:
        raise SowExtractionError(
            "No AI provider is configured -- set GEMINI_API_KEY, ANTHROPIC_API_KEY, or the Azure OpenAI vars in the backend .env file."
        )

    prompt = _build_extraction_prompt(sow_text)
    messages = [{"role": "user", "content": prompt}]

    for provider in providers:
        try:
            turn = provider.generate_with_tools(messages, [], max_tokens=2000)
        except Exception:
            continue
        content = (turn or {}).get("content")
        if not content:
            continue
        cleaned = content.strip()
        if cleaned.startswith("```"):
            cleaned = cleaned.strip("`")
            cleaned = cleaned[4:] if cleaned.lower().startswith("json") else cleaned
        try:
            parsed = json.loads(cleaned.strip())
        except json.JSONDecodeError:
            continue
        if isinstance(parsed.get("roles_required"), list):
            return parsed

    raise SowExtractionError("The AI could not extract a valid result from this SOW -- try again or check the file has real, readable text.")


def _load_extractions_index() -> dict:
    if not EXTRACTIONS_JSON_PATH.exists():
        return {}
    try:
        return json.loads(EXTRACTIONS_JSON_PATH.read_text())
    except json.JSONDecodeError:
        return {}


def _save_extractions_index(index: dict) -> None:
    EXTRACTIONS_JSON_PATH.write_text(json.dumps(index, indent=2))


def _index_key(project_code: str, filename: str) -> str:
    return f"{project_code}::{filename}"


def get_cached_extraction(project_code: str, filename: str) -> dict | None:
    return _load_extractions_index().get(_index_key(project_code, filename))


def extract_sow_requirements(project_code: str, filename: str) -> dict:
    path = get_sow_file_path(project_code, filename)
    if path is None:
        raise SowExtractionError(f"No SOW file '{filename}' on record for project '{project_code}'.")

    content = Path(path).read_bytes()
    sow_text = _extract_text(filename, content)
    result = _call_llm_for_extraction(sow_text)

    real_designations_lower = {d.lower() for d in _real_designations()}
    roles = []
    for r in result.get("roles_required", []):
        if not isinstance(r, dict):
            continue
        role_text = str(r.get("role_text") or "").strip()
        if not role_text:
            continue
        try:
            count = max(1, int(r.get("count", 1)))
        except (TypeError, ValueError):
            count = 1
        roles.append({
            "role_text": role_text,
            "count": count,
            "named_person": r.get("named_person") or None,
            "matches_real_designation": role_text.lower() in real_designations_lower,
        })

    record = {
        "project_code": project_code,
        "filename": filename,
        "extracted_at": pd.Timestamp.now().isoformat(),
        "roles_required": roles,
        "skills_required": [str(s).strip() for s in result.get("skills_required", []) if str(s).strip()],
        "competency_areas_required": [str(c).strip() for c in result.get("competency_areas_required", []) if str(c).strip()],
        "client_name": result.get("client_name") or None,
        "project_reference": result.get("project_reference") or None,
        "engagement_duration": result.get("engagement_duration") or None,
        "scope_summary": result.get("scope_summary") or None,
    }

    index = _load_extractions_index()
    index[_index_key(project_code, filename)] = record
    _save_extractions_index(index)
    return record
